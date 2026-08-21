import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def add_heading_styled(doc, text, level):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = True
    
    if level == 1:
        run.font.size = Pt(17)
        run.font.color.rgb = RGBColor(0x0F, 0x3B, 0x68) # Navy Blue
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1B, 0x6B, 0x93) # Teal Blue
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x28, 0x52, 0x7A)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
    return p

def add_paragraph_styled(doc, text="", bold_prefix="", italic=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        r_prefix.font.name = 'Calibri'
        r_prefix.font.size = Pt(10.5)
        r_prefix.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        r_prefix.bold = True
        
    if text:
        r_text = p.add_run(text)
        r_text.font.name = 'Calibri'
        r_text.font.size = Pt(10.5)
        r_text.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        r_text.italic = italic
        
    return p

def build_docx_en(output_path, img_path):
    doc = Document()
    
    # Page setup (A4 standard)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Title Cover Block
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(16)
    title_p.paragraph_format.space_after = Pt(4)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = title_p.add_run("Statement of Work (SOW) Specification")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(20)
    r_title.font.color.rgb = RGBColor(0x0F, 0x3B, 0x68)
    r_title.bold = True
    
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(14)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = sub_p.add_run("PMT Platform: 12-Step Professional Project Management Pipeline")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = RGBColor(0x1B, 0x6B, 0x93)
    r_sub.bold = True

    # Meta Table (Document Info)
    info_tbl = doc.add_table(rows=4, cols=2)
    info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(info_tbl, "B0C4DE")
    col_widths = [Inches(2.2), Inches(4.6)]
    
    doc_info = [
        ("Project / System Name:", "PMT (Project Management & Execution Platform)"),
        ("Prepared by:", "Project Management (PM) & Solutions Architecture (SA) Team"),
        ("Document Objective:", "Scope of Work & Technical Specification Attachment for SOW / RFP"),
        ("Document Version & Date:", "Version 1.0 (Official Release)")
    ]
    
    for row_idx, (k, v) in enumerate(doc_info):
        row = info_tbl.rows[row_idx]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width, c1.width = col_widths[0], col_widths[1]
        set_cell_background(c0, "F0F4F8")
        set_cell_background(c1, "FAFCFE")
        set_cell_margins(c0, 60, 60, 100, 100)
        set_cell_margins(c1, 60, 60, 100, 100)
        
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(0)
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(k)
        r0.font.name = 'Calibri'
        r0.font.size = Pt(10)
        r0.bold = True
        r0.font.color.rgb = RGBColor(0x0F, 0x3B, 0x68)
        
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(v)
        r1.font.name = 'Calibri'
        r1.font.size = Pt(10)
        r1.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    # Embed Image if exists
    if os.path.exists(img_path):
        img_p = doc.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.paragraph_format.space_before = Pt(8)
        img_p.paragraph_format.space_after = Pt(4)
        img_run = img_p.add_run()
        img_run.add_picture(img_path, width=Inches(6.8))
        
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_after = Pt(12)
        cap_r = cap_p.add_run("Figure 1: PMT 12-Step Project Pipeline Architecture & Role Overview")
        cap_r.font.name = 'Calibri'
        cap_r.font.size = Pt(9.5)
        cap_r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        cap_r.italic = True

    # 1. Core User Roles Matrix
    add_heading_styled(doc, "1. Core User Roles & Access Matrix", 1)
    add_paragraph_styled(doc, "The PMT Platform incorporates strict Separation of Duties (SoD) mapped to four fundamental enterprise roles:")

    role_tbl = doc.add_table(rows=5, cols=3)
    role_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(role_tbl, "CCCCCC")
    role_widths = [Inches(1.6), Inches(1.2), Inches(4.0)]
    
    headers = ["Role Name", "Role Code", "Functional Scope & System Responsibilities"]
    for i, h in enumerate(headers):
        cell = role_tbl.rows[0].cells[i]
        cell.width = role_widths[i]
        set_cell_background(cell, "0F3B68")
        set_cell_margins(cell, 80, 80, 100, 100)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.name = 'Calibri'
        r.font.size = Pt(10.5)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    roles_data = [
        ("Admin (System Administrator)", "ROLE_ADMIN", "Full administrative control over Master Data (Project Types, Zones, Templates), Service Price Book configurations, base labor cost rates, and company-wide user permissions."),
        ("Manager / PM (Project Manager)", "ROLE_PM", "Responsible for end-to-end project planning, Gantt Baseline creation, work delegation, BOQ & Quotation issuance, timesheet approvals, and project budget control."),
        ("Employee / Technician (Field Team)", "ROLE_TECH", "Operates mobile web interface for GPS site check-in, captures live photo proof, submits daily task-based timesheets, and updates card progression on Kanban boards."),
        ("Executive (Management)", "ROLE_EXEC", "Accesses real-time Executive Dashboards to review portfolio-wide project health, budget burn rates, Gross Margin (P&L), and schedule drift indices.")
    ]
    
    for row_idx, (r_name, r_code, r_desc) in enumerate(roles_data, start=1):
        row = role_tbl.rows[row_idx]
        bg = "F9FBFD" if row_idx % 2 == 1 else "FFFFFF"
        for c_idx, text in enumerate([r_name, r_code, r_desc]):
            cell = row.cells[c_idx]
            cell.width = role_widths[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, 60, 60, 80, 80)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = 'Calibri'
            r.font.size = Pt(9.5)
            if c_idx == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(0x0F, 0x3B, 0x68)
            else:
                r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 2. Key Quality Control Features
    add_heading_styled(doc, "2. Key Quality Control Pillars", 1)
    
    features = [
        ("GPS & Camera Proof of Work:", " Enforces true physical site presence via satellite Geolocation API and mandates instant live photo capture via HTML5 Camera API to eliminate proxy attendance and falsified timesheets."),
        ("Gantt Baseline & Drift Variance Engine:", " Locks initial project schedule as a Frozen Baseline. Automatically computes variance against real-time live execution to quantify delay days and prevent budget overrun."),
        ("Agile Kanban Board Integration:", " Provides intuitive task-based board workflow linked directly to labor timesheets. Team members transition cards across To Do, In Progress, and Done with automated cost accruals."),
        ("Real-time Executive Dashboard & P&L:", " Aggregates key portfolio health metrics including Budget Burn Rate, Cash Flow Milestones, and proactive risk indicators for data-driven strategic decisions.")
    ]
    for title, desc in features:
        add_paragraph_styled(doc, text=desc, bold_prefix=title, space_after=3)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 3. 12-Step Pipeline Specification
    add_heading_styled(doc, "3. The 12-Step Project Pipeline Specification", 1)
    add_paragraph_styled(doc, "The PMT project lifecycle is partitioned into 3 consecutive phases across 12 operational steps:")

    steps_data_en = [
        # Phase 1
        {
            "num": 1,
            "phase": "Phase 1: Pre-Construction",
            "name": "Lead & Requirement Gathering",
            "obj": "Capture and qualify target customer leads, contact details, project categorization (Renovation, Built-in, System Install), scope requirements, and initial budget envelope.",
            "actor": "Responsible: Sales / PM | Accountable: PM | Informed: Admin",
            "inputs": "Customer name, contact channels (Phone/LINE/Email), site geolocation pin, estimated budget, service type.",
            "process": "1. Sales or PM inputs incoming inquiries into the CRM/Lead module.\n2. Create new Lead with Google Maps location pin and assign priority rating.\n3. System generates unique Lead Code and dispatches automated notifications to designated regional PMs.",
            "sa": "Entity `leads` stores records with auto-generated alphanumeric code (e.g., LD-2026-0001) and performs GPS/contact data validation.",
            "output": "Lead record created with 'New Lead' status, visible on the Central Leads Management Board.",
            "sow": "The platform must support real-time lead capture, search/filtering, and accurate geospatial mapping."
        },
        {
            "num": 2,
            "phase": "Phase 1: Pre-Construction",
            "name": "Survey Booking & Scheduling",
            "obj": "Schedule and coordinate field surveyor site visits with automated time-slot verification to prevent double-booking and synchronize customer appointments.",
            "actor": "Responsible: PM / Booking Officer | Accountable: PM | Informed: Field Surveyor, Customer",
            "inputs": "Qualified Lead ID, customer preferred appointment date/time slot, available surveyor technician roster.",
            "process": "1. PM checks master resource availability via the Central Booking Calendar.\n2. Allocate appointment slot and assign field technician.\n3. System issues automated confirmation notifications via SMS/Email/LINE to customer and technician.",
            "sa": "Entity `surveys` incorporates a Conflict Detection Algorithm preventing overlapping bookings and updates Lead status to 'Survey Scheduled'.",
            "output": "Confirmed survey appointment logged in Central Calendar and synchronized to technician's mobile view.",
            "sow": "Automated resource conflict detection and multi-channel appointment notification dispatch."
        },
        {
            "num": 3,
            "phase": "Phase 1: Pre-Construction",
            "name": "Survey QC Inspection & As-is Evidence",
            "obj": "Inspect physical site conditions, measure dimensions, assess structural/utility constraints, and capture pre-existing condition photos (Pre-work Proof) to mitigate liability disputes.",
            "actor": "Responsible: Field Surveyor / Inspector | Accountable: PM | Consulted: Customer",
            "inputs": "Survey appointment record, Digital Survey Checklist template, mobile device for dimension logging and photo capture.",
            "process": "1. Surveyor arrives at site and executes Mobile Check-in.\n2. Fills digital checklist (Width x Length x Height, surface conditions, electrical/plumbing status).\n3. Captures live on-site condition photos via integrated camera API.\n4. Customer or site representative signs off digitally on the mobile screen.",
            "sa": "Invokes HTML5 Geolocation API and Camera API to stamp GPS coordinates and timestamp metadata onto photo artifacts; compiles Digital Survey Report.",
            "output": "Digital Survey Report with stamped photos and digital signature; status updated to 'Survey Completed'.",
            "sow": "Mandatory live camera capture with GPS metadata stamping and digital customer acknowledgment."
        },
        {
            "num": 4,
            "phase": "Phase 1: Pre-Construction",
            "name": "2D/3D Design & Space Planning",
            "obj": "Develop and upload 2D architectural layouts and 3D realistic renderings for customer conceptualization and formal design sign-off prior to cost estimation.",
            "actor": "Responsible: Architect / Designer | Accountable: PM | Consulted: Customer",
            "inputs": "Survey report measurements and photos, architectural guidelines, customer styling preferences.",
            "process": "1. Designer drafts 2D floor plans and renders 3D perspectives based on survey data.\n2. Uploads design packages and drawings to the project repository.\n3. PM presents design concepts to the customer for review.\n4. Customer approves design; system logs formal Design Approval timestamp.",
            "sa": "Entity `project_designs` with built-in Revision Control (Rev A, B, C) preventing obsolete drawing usage and enabling in-browser PDF/Image viewing.",
            "output": "Approved 2D/3D design package; project status progresses to 'Design Approved'.",
            "sow": "Design version control repository with audit trail of customer approval records."
        },
        
        # Phase 2
        {
            "num": 5,
            "phase": "Phase 2: Construction (Budget & Site Launch)",
            "name": "BOQ Cost Estimation & Pricing",
            "obj": "Generate accurate Bill of Quantities (BOQ) covering materials, labor, and overheads dynamically pulled from the Service Price Book to ensure target Gross Profit Margins.",
            "actor": "Responsible: PM / Quantity Surveyor (QS) | Accountable: PM / Executive",
            "inputs": "Approved design drawings, Service Price Book Master Data, itemized work quantities.",
            "process": "1. PM/QS generates a BOQ document linked to the approved design.\n2. Selects standardized items from the Service Price Book (auto-populates cost rate and benchmark sell rate).\n3. Enters quantities; system calculates total estimated costs and recommended price.\n4. Evaluates Gross Margin % and establishes the Baseline Cost Budget.",
            "sa": "Entities `boq_items` and `service_price_book` execute formula: Material + Labor + Overhead + Margin = Total Selling Price with locked calculation logic.",
            "output": "Itemized BOQ document and frozen Baseline Project Budget.",
            "sow": "Standardized Price Book integration with automated margin calculations and locked baseline budget creation."
        },
        {
            "num": 6,
            "phase": "Phase 2: Construction (Budget & Site Launch)",
            "name": "Quotation Approval & Contracting",
            "obj": "Convert BOQ into an official Quotation, present payment milestone terms to the customer, and secure formal contract agreement / purchase order.",
            "actor": "Responsible: PM / Sales | Accountable: PM | Approver: Authorized Signatory / Customer",
            "inputs": "Validated BOQ document, payment milestone schedule, contract terms and conditions.",
            "process": "1. System generates official Quotation with formal numbering and payment schedule.\n2. Digital quotation link dispatched to customer.\n3. Customer confirms and signs via E-Signature or uploads approved Purchase Order (PO).\n4. PM approves in system; Lead is automatically converted into an 'Active Project' with Gantt and Kanban structures.",
            "sa": "Entities `quotations`, `projects`, and `contracts` handle entity transformation from Lead to Active Project and auto-instantiate Milestone WBS templates.",
            "output": "Signed Contract / Approved Quotation; Active Project instance initialized with status 'Quotation Approved'.",
            "sow": "Automated quotation numbering, contract repository, and instant project instance generation upon approval."
        },
        {
            "num": 7,
            "phase": "Phase 2: Construction (Budget & Site Launch)",
            "name": "Down Payment & Financial Entry",
            "obj": "Record receipt of initial advance deposit (Down Payment) prior to physical site mobilization to mitigate organizational financial exposure.",
            "actor": "Responsible: Finance / Accounting | Accountable: PM | Informed: Customer",
            "inputs": "Approved Quotation, customer payment transfer slip / bank confirmation.",
            "process": "1. Customer executes deposit payment as per contract terms.\n2. Finance verifies funds and logs payment transaction in system with attached slip.\n3. Milestone 1 financial status changes to 'Paid'.\n4. System unlocks authorization for field mobilization and Site Check-In.",
            "sa": "Entities `project_payments` and `invoices` log Cash Inflow and dispatch authorization signals to the Project Execution module.",
            "output": "Official Payment Receipt; project status updated to 'Ready for Site Entry'.",
            "sow": "Strict financial gatekeeper preventing site mobilization prior to validated down payment confirmation."
        },
        {
            "num": 8,
            "phase": "Phase 2: Construction (Budget & Site Launch)",
            "name": "Site Check-In & Proof of Presence",
            "obj": "Enforce physical site attendance verification for field teams via satellite Geofencing and live photo proof to prevent false check-ins.",
            "actor": "Responsible: Lead Technician / Field Team | Accountable: PM | Informed: Customer",
            "inputs": "Technician mobile web application, GPS device sensors, built-in camera hardware.",
            "process": "1. Field team arrives at construction site, opens mobile interface, and selects active project.\n2. Clicks 'Site Check-In'; system evaluates device GPS against registered site coordinates (Geofence radius 200-500m).\n3. Technician captures live selfie with site background via live camera.\n4. Confirms check-in; real-time attendance alert appears on PM Dashboard.",
            "sa": "Evaluates Haversine distance algorithm against site coordinates; restricts input to `capture=camera` and persists to `site_checkins`.",
            "output": "Verified Site Check-in log with coordinates, timestamp, and photo displayed on Live Project Dashboard.",
            "sow": "Automated GPS geofence validation and live photo capture requirement with out-of-range warnings."
        },
        
        # Phase 3
        {
            "num": 9,
            "phase": "Phase 3: Completion & Review",
            "name": "Execution & Daily Timesheet Logging",
            "obj": "Execute construction tasks according to schedule, log worker labor hours (Timesheets), and update Agile Kanban cards for automated labor cost accruals.",
            "actor": "Responsible: Field Technicians / Workers | Accountable: PM | Consulted: Site Foreman",
            "inputs": "Kanban task cards, actual regular and overtime work hours, daily work progress photos.",
            "process": "1. Technicians carry out assigned tasks and transition Kanban cards (To Do -> In Progress -> Done).\n2. At end of shift, workers log hours and attach daily completion photos.\n3. Submit daily timesheet batch to PM for formal review and approval.",
            "sa": "Entities `timesheets`, `tasks`, and `kanban_cards` multiply logged hours by user hourly cost rates to accrue Actual Labor Costs in real-time.",
            "output": "Daily Timesheet entries logged and cumulative labor expenses updated in project ledger.",
            "sow": "Direct integration between Kanban tasks and timesheets with automatic labor cost rate multiplication."
        },
        {
            "num": 10,
            "phase": "Phase 3: Completion & Review",
            "name": "Daily QC & Progress Update (Baseline Drift)",
            "obj": "Perform milestone quality inspections, approve submitted timesheets, and record % actual progress against the Gantt Baseline to quantify schedule drift.",
            "actor": "Responsible: PM / Site Engineer / QC Inspector | Accountable: PM",
            "inputs": "Pending timesheet submissions, daily progress photos, milestone QC checklist criteria.",
            "process": "1. PM/Engineer inspects site deliverables physically or via digital photo submissions.\n2. Reviews and Approves or Rejects submitted timesheets with explanatory notes.\n3. Evaluates QC checklist criteria for the active milestone.\n4. Updates % actual progress; system calculates schedule variance and delay drift days.",
            "sa": "Gantt Engine computes Baseline vs Live Timeline variance; stores inspection audits in `qc_inspections` and updates EVM metrics.",
            "output": "Approved Timesheets, Daily QC Inspection Report, and updated Gantt Drift visualization.",
            "sow": "Formal PM timesheet approval workflow and real-time Gantt Baseline vs Actual Drift variance analysis."
        },
        {
            "num": 11,
            "phase": "Phase 3: Completion & Review",
            "name": "Final QC & Customer Handover",
            "obj": "Conduct comprehensive final inspection, resolve punch list defect items, and execute official digital handover sign-off with the customer.",
            "actor": "Responsible: PM / QC Lead | Accountable: PM | Approver: Customer",
            "inputs": "100% completed project deliverables, Defect Punch List, Handover Certificate template.",
            "process": "1. PM conducts final walk-through inspection with customer.\n2. Logs any punch list defect items in the system and assigns rectification tasks.\n3. Upon 100% defect clearance, customer executes Digital Handover Sign-off on screen.\n4. Record final milestone payment receipt and initiate warranty validity period.",
            "sa": "Entities `handover_certificates` and `defect_items` compile official PDF Handover Certificate with digital signature and update status to 'Delivered / Pending Closeout'.",
            "output": "Signed Handover Certificate, cleared defect logs, and activated warranty certificate.",
            "sow": "Digital punch list defect management lifecycle and electronic customer handover sign-off."
        },
        {
            "num": 12,
            "phase": "Phase 3: Completion & Review",
            "name": "Project Closeout & P&L Evaluation",
            "obj": "Formally close project, compute actual Profit & Loss (Revenue vs Actual Materials + Timesheet Labor + Expenses), and capture Customer Satisfaction (CSAT) scores.",
            "actor": "Responsible: PM / Finance | Accountable: Executive | Informed: All Stakeholders",
            "inputs": "Fully settled revenue invoices, all actual expense bills and approved labor timesheets, completed CSAT survey.",
            "process": "1. Finance and PM audit all final revenues and accrued costs.\n2. System computes final Gross Profit Margin %, net profitability, and final schedule variance.\n3. Customer completes digital CSAT evaluation.\n4. PM clicks 'Close Project'; system archives project records into permanent read-only state.",
            "sa": "Entity `projects` sets status to 'Closed', aggregates project KPIs into Executive Data Lake, and locks historical records against retrospective edits.",
            "output": "Final Project Closeout P&L Report, CSAT audit score, and archived project record.",
            "sow": "Automated financial P&L reconciliation comparing budget vs actuals and permanent project lock upon closeout."
        }
    ]

    current_phase = ""
    for step in steps_data_en:
        if step["phase"] != current_phase:
            current_phase = step["phase"]
            add_heading_styled(doc, current_phase, 2)
            
        add_heading_styled(doc, f"Step {step['num']}: {step['name']}", 3)
        
        # Step Table Specification
        tbl = doc.add_table(rows=7, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl, "B0C4DE")
        s_widths = [Inches(1.8), Inches(5.0)]
        
        fields = [
            ("Business Objectives", step["obj"]),
            ("RACI Roles", step["actor"]),
            ("Inputs / Prerequisites", step["inputs"]),
            ("Process Workflow", step["process"]),
            ("System & SA Mechanics", step["sa"]),
            ("Outputs / Deliverables", step["output"]),
            ("SOW Acceptance Criteria", step["sow"])
        ]
        
        for r_idx, (f_label, f_val) in enumerate(fields):
            row = tbl.rows[r_idx]
            c0, c1 = row.cells[0], row.cells[1]
            c0.width, c1.width = s_widths[0], s_widths[1]
            set_cell_background(c0, "F0F4F8")
            set_cell_background(c1, "FFFFFF" if r_idx % 2 == 0 else "FAFCFE")
            set_cell_margins(c0, 50, 50, 80, 80)
            set_cell_margins(c1, 50, 50, 80, 80)
            
            p0 = c0.paragraphs[0]
            p0.paragraph_format.space_before = Pt(0)
            p0.paragraph_format.space_after = Pt(0)
            r0 = p0.add_run(f_label)
            r0.font.name = 'Calibri'
            r0.font.size = Pt(9.5)
            r0.bold = True
            r0.font.color.rgb = RGBColor(0x0F, 0x3B, 0x68)
            
            p1 = c1.paragraphs[0]
            p1.paragraph_format.space_before = Pt(0)
            p1.paragraph_format.space_after = Pt(0)
            r1 = p1.add_run(f_val)
            r1.font.name = 'Calibri'
            r1.font.size = Pt(9.5)
            r1.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 4. RACI Responsibility Matrix
    add_heading_styled(doc, "4. RACI Responsibility Matrix Across Pipeline Steps", 1)
    add_paragraph_styled(doc, "Legend: R = Responsible (Operator), A = Accountable (Approver/Owner), C = Consulted (Advisor), I = Informed (Observer)")

    raci_tbl = doc.add_table(rows=13, cols=6)
    raci_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(raci_tbl, "CCCCCC")
    raci_widths = [Inches(0.6), Inches(2.6), Inches(0.9), Inches(0.9), Inches(0.9), Inches(0.9)]
    
    raci_headers = ["#", "Pipeline Stage / Step", "Admin", "PM", "Technician", "Executive"]
    for i, h in enumerate(raci_headers):
        cell = raci_tbl.rows[0].cells[i]
        cell.width = raci_widths[i]
        set_cell_background(cell, "0F3B68")
        set_cell_margins(cell, 60, 60, 60, 60)
        p = cell.paragraphs[0]
        if i >= 2:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    raci_matrix = [
        ("1", "Lead & Requirement Gathering", "I", "R/A", "-", "I"),
        ("2", "Survey Booking & Scheduling", "I", "R/A", "C", "I"),
        ("3", "Survey QC Inspection", "-", "A", "R", "I"),
        ("4", "2D/3D Design & Planning", "-", "A", "-", "I"),
        ("5", "BOQ Cost Estimation", "C", "R/A", "-", "I"),
        ("6", "Quotation Approval & Contract", "I", "R/A", "-", "A/I"),
        ("7", "Down Payment Entry", "I", "R/A", "-", "I"),
        ("8", "Site Check-In (GPS & Photo)", "-", "A", "R", "I"),
        ("9", "Execution & Daily Timesheet", "-", "A", "R", "I"),
        ("10", "Daily QC & Progress Update", "-", "R/A", "C", "I"),
        ("11", "Final QC & Customer Handover", "-", "R/A", "C", "I"),
        ("12", "Project Closeout & P&L", "I", "R", "-", "A")
    ]
    
    for row_idx, r_data in enumerate(raci_matrix, start=1):
        row = raci_tbl.rows[row_idx]
        bg = "F9FBFD" if row_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(r_data):
            cell = row.cells[c_idx]
            cell.width = raci_widths[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, 50, 50, 60, 60)
            p = cell.paragraphs[0]
            if c_idx >= 2 or c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            r.font.name = 'Calibri'
            r.font.size = Pt(9.5)
            if "R" in val or "A" in val:
                r.bold = True
                r.font.color.rgb = RGBColor(0x0F, 0x3B, 0x68)
            else:
                r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 5. Sign-off Section
    add_heading_styled(doc, "5. Statement of Work (SOW) Acceptance Sign-Off", 1)
    add_paragraph_styled(doc, "This PMT SOW Specification has been thoroughly reviewed and agreed upon by all principal stakeholders as the binding functional and technical baseline for implementation and quality acceptance:")

    sign_tbl = doc.add_table(rows=2, cols=2)
    sign_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(sign_tbl, "CCCCCC")
    sign_widths = [Inches(3.4), Inches(3.4)]
    
    sign_boxes = [
        ("For: Client / Project Executive Authority", "\n\n\nSignature: ........................................................\nName: ( ........................................................ )\nTitle: ....................................................\nDate: ......./......./............"),
        ("For: Project Manager & Solutions Architecture Lead", "\n\n\nSignature: ........................................................\nName: ( ........................................................ )\nTitle: ....................................................\nDate: ......./......./............")
    ]
    
    for c_idx, (s_title, s_content) in enumerate(sign_boxes):
        cell = sign_tbl.rows[0].cells[c_idx]
        cell.width = sign_widths[c_idx]
        set_cell_background(cell, "F0F4F8")
        set_cell_margins(cell, 60, 60, 80, 80)
        p = cell.paragraphs[0]
        r = p.add_run(s_title)
        r.font.name = 'Calibri'
        r.font.size = Pt(9.5)
        r.bold = True
        r.font.color.rgb = RGBColor(0x0F, 0x3B, 0x68)
        
        cell2 = sign_tbl.rows[1].cells[c_idx]
        cell2.width = sign_widths[c_idx]
        set_cell_background(cell2, "FFFFFF")
        set_cell_margins(cell2, 80, 80, 80, 80)
        p2 = cell2.paragraphs[0]
        r2 = p2.add_run(s_content)
        r2.font.name = 'Calibri'
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.save(output_path)
    print(f"Successfully generated: {output_path}")

if __name__ == "__main__":
    out_file = r"c:\atgv\vbooking\PMT_12_Steps_SOW_Specification_EN.docx"
    img_file = r"C:\Users\isara\.gemini\antigravity\brain\a47b4d83-2349-4a12-bd2a-ed2291707780\.user_uploaded\media_1787020048617.jpg"
    build_docx_en(out_file, img_file)
