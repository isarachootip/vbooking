import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def markdown_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Base Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    def flush_table(tbl_lines):
        if not tbl_lines:
            return
        parsed_rows = []
        for l in tbl_lines:
            l = l.strip()
            if l.startswith('|') and l.endswith('|'):
                l = l[1:-1]
            cells = [c.strip() for c in l.split('|')]
            if len(cells) > 0 and not all(re.match(r'^:?-+:?$', c) for c in cells):
                parsed_rows.append(cells)

        if not parsed_rows:
            return

        num_cols = max(len(r) for r in parsed_rows)
        table = doc.add_table(rows=len(parsed_rows), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        for row_idx, row_data in enumerate(parsed_rows):
            for col_idx in range(num_cols):
                val = row_data[col_idx] if col_idx < len(row_data) else ""
                cell = table.cell(row_idx, col_idx)
                set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                
                # Format cell text
                # strip bold tags for display or add bold runs
                clean_text = val.replace('**', '')
                run = p.add_run(clean_text)
                run.font.size = Pt(9.5)
                run.font.name = 'Calibri'

                if row_idx == 0:
                    set_cell_background(cell, '1E293B')  # dark slate blue
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                else:
                    if row_idx % 2 == 1:
                        set_cell_background(cell, 'F8FAFC')
                    else:
                        set_cell_background(cell, 'FFFFFF')
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def flush_code(c_lines):
        if not c_lines:
            return
        box_table = doc.add_table(rows=1, cols=1)
        box_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = box_table.cell(0, 0)
        set_cell_background(cell, '0F172A')
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        
        run = p.add_run("\n".join(c_lines))
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x38, 0xBD, 0xF8)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    for line in lines:
        raw_line = line.rstrip('\r\n')
        stripped = raw_line.strip()

        # Code block handling
        if stripped.startswith('```'):
            if in_code_block:
                flush_code(code_lines)
                code_lines = []
                in_code_block = False
            else:
                if in_table:
                    flush_table(table_lines)
                    table_lines = []
                    in_table = False
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(raw_line)
            continue

        # Table handling
        if stripped.startswith('|') and '|' in stripped[1:]:
            in_table = True
            table_lines.append(stripped)
            continue
        elif in_table:
            flush_table(table_lines)
            table_lines = []
            in_table = False

        if not stripped:
            continue

        # Heading 1
        if stripped.startswith('# '):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(20)
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Dark Navy
            continue

        # Heading 2
        if stripped.startswith('## '):
            text = stripped[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(15)
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(0x02, 0x84, 0xC7) # Bright Blue
            continue

        # Heading 3
        if stripped.startswith('### '):
            text = stripped[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12.5)
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            continue

        # Heading 4
        if stripped.startswith('#### '):
            text = stripped[5:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11.5)
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
            continue

        # Divider
        if stripped in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run('—' * 45)
            r.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
            continue

        # List items (bullets)
        if stripped.startswith('* ') or stripped.startswith('- ') or stripped.startswith('  * ') or stripped.startswith('    * '):
            indent_level = len(raw_line) - len(raw_line.lstrip())
            content = stripped.lstrip('*- ').strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1.5)
            p.paragraph_format.space_after = Pt(1.5)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.left_indent = Inches(0.25 * (indent_level // 2 + 1))
            
            # Simple bold parser
            parts = re.split(r'(\*\*.*?\*\*)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                    r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
                else:
                    # remove backticks
                    clean = part.replace('`', '')
                    r = p.add_run(clean)
                    r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            continue

        # Numbered list
        num_match = re.match(r'^(\d+\.)\s+(.*)$', stripped)
        if num_match:
            prefix, content = num_match.group(1), num_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.left_indent = Inches(0.25)
            
            r_num = p.add_run(prefix + " ")
            r_num.bold = True
            r_num.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
            
            parts = re.split(r'(\*\*.*?\*\*)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    clean = part.replace('`', '')
                    r = p.add_run(clean)
            continue

        # Regular Paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.2
        
        parts = re.split(r'(\*\*.*?\*\*)', stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                r = p.add_run(part[2:-2])
                r.bold = True
            elif part.startswith('*') and part.endswith('*'):
                r = p.add_run(part[1:-1])
                r.italic = True
                r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            else:
                clean = part.replace('`', '')
                r = p.add_run(clean)

    if in_table:
        flush_table(table_lines)

    doc.save(docx_path)
    print(f"Successfully converted to {docx_path}")

if __name__ == '__main__':
    md_file = 'SYSTEM_TRAINING_MANUAL.md'
    docx_file = 'NexTime_System_Training_Manual.docx'
    markdown_to_docx(md_file, docx_file)
