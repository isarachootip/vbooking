import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def add_heading_styled(doc, text, level):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.font.name = 'TH Sarabun New'
    run.bold = True
    
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x0F, 0x3B, 0x68) # Navy Blue
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1B, 0x6B, 0x93) # Teal Blue
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    elif level == 3:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x28, 0x52, 0x7A)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    return p

def add_paragraph_styled(doc, text="", bold_prefix="", italic=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        r_prefix.font.name = 'TH Sarabun New'
        r_prefix.font.size = Pt(13)
        r_prefix.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        r_prefix.bold = True
        
    if text:
        r_text = p.add_run(text)
        r_text.font.name = 'TH Sarabun New'
        r_text.font.size = Pt(13)
        r_text.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        r_text.italic = italic
        
    return p

def build_docx(output_path, img_path):
    doc = Document()
    
    # Page setup (A4 standard)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Title Cover Block
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(20)
    title_p.paragraph_format.space_after = Pt(4)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = title_p.add_run("เอกสารข้อกำหนดขอบเขตงานระบบ (Statement of Work - SOW)")
    r_title.font.name = 'TH Sarabun New'
    r_title.font.size = Pt(22)
    r_title.font.color.rgb = RGBColor(0x0F, 0x3B, 0x68)
    r_title.bold = True
    
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(18)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = sub_p.add_run("ระบบ PMT: 12 ขั้นตอนการบริหารโครงการอย่างมืออาชีพ (12-Step Project Pipeline Specification)")
    r_sub.font.name = 'TH Sarabun New'
    r_sub.font.size = Pt(16)
    r_sub.font.color.rgb = RGBColor(0x1B, 0x6B, 0x93)
    r_sub.bold = True

    # Meta Table (Document Info)
    info_tbl = doc.add_table(rows=4, cols=2)
    info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(info_tbl, "B0C4DE")
    col_widths = [Inches(2.2), Inches(4.6)]
    
    doc_info = [
        ("ชื่อโครงการ / ระบบ:", "PMT (Project Management & Execution Platform)"),
        ("จัดทำโดย (Prepared by):", "ทีม Project Manager (PM) & Solutions Architect (SA)"),
        ("วัตถุประสงค์ (Purpose):", "เอกสารประกอบสัญญาจ้างและข้อกำหนดทางเทคนิค (SOW / TOR)"),
        ("เวอร์ชันเอกสาร / วันที่:", "Version 1.0 (Official Release)")
    ]
    
    for row_idx, (k, v) in enumerate(doc_info):
        row = info_tbl.rows[row_idx]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width, c1.width = col_widths[0], col_widths[1]
        set_cell_background(c0, "F0F4F8")
        set_cell_background(c1, "FAFCFE")
        set_cell_margins(c0, 80, 80, 120, 120)
        set_cell_margins(c1, 80, 80, 120, 120)
        
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(0)
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(k)
        r0.font.name = 'TH Sarabun New'
        r0.font.size = Pt(12)
        r0.bold = True
        r0.font.color.rgb = RGBColor(0x0F, 0x3B, 0x68)
        
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(v)
        r1.font.name = 'TH Sarabun New'
        r1.font.size = Pt(12)
        r1.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    # Embed Image if exists
    if os.path.exists(img_path):
        img_p = doc.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.paragraph_format.space_before = Pt(10)
        img_p.paragraph_format.space_after = Pt(6)
        img_run = img_p.add_run()
        img_run.add_picture(img_path, width=Inches(6.8))
        
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_after = Pt(16)
        cap_r = cap_p.add_run("ภาพที่ 1: แผนผังวงจรการทำงาน 12 ขั้นตอนและโครงสร้างระบบ PMT")
        cap_r.font.name = 'TH Sarabun New'
        cap_r.font.size = Pt(11)
        cap_r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        cap_r.italic = True

    # 1. บทบาทหลักในระบบ (Core User Roles)
    add_heading_styled(doc, "1. โครงสร้างบทบาทผู้ใช้งานในระบบ (Core User Roles Matrix)", 1)
    add_paragraph_styled(doc, "ระบบ PMT ได้รับการออกแบบให้สอดคล้องกับมาตรฐานการทำงานขององค์กร โดยแบ่งแยกหน้าที่และความรับผิดชอบ (Separation of Duties) ออกเป็น 4 บทบาทหลัก ดังนี้:")

    role_tbl = doc.add_table(rows=5, cols=3)
    role_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(role_tbl, "CCCCCC")
    role_widths = [Inches(1.5), Inches(1.3), Inches(4.0)]
    
    headers = ["บทบาท (Role)", "รหัสสิทธิ์", "ขอบเขตหน้าที่และความรับผิดชอบในระบบ"]
    for i, h in enumerate(headers):
        cell = role_tbl.rows[0].cells[i]
        cell.width = role_widths[i]
        set_cell_background(cell, "0F3B68")
        set_cell_margins(cell, 100, 100, 120, 120)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.name = 'TH Sarabun New'
        r.font.size = Pt(13)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    roles_data = [
        ("Admin (ผู้ดูแลระบบ)", "ROLE_ADMIN", "มีสิทธิ์สูงสุดในการจัดการข้อมูลหลัก (Master Data), กำหนดอัตราค่าแรง (Cost Rates), ตั้งค่า Service Price Book, และควบคุมสิทธิ์การใช้งานของพนักงานทุกคน"),
        ("Manager / PM (ผู้จัดการโครงการ)", "ROLE_PM", "รับผิดชอบการบริหารโครงการ, วางแผนงาน Gantt Baseline, มอบหมายงานให้ทีมช่าง, จัดทำใบเสนอราคา/BOQ, อนุมัติ Timesheet และควบคุมต้นทุนไม่ให้เกินงบ"),
        ("Employee / Technician (ทีมช่างปฏิบัติการ)", "ROLE_TECH", "ใช้งานผ่านมือถือ/แท็บเล็ตเพื่อเช็คอินพิกัด GPS ณ หน้างานจริง, บันทึกชั่วโมงการทำงาน (Timesheet) พร้อมรูปถ่ายหลักฐาน และอัปเดตความคืบหน้ารายวันบนบอร์ด Kanban"),
        ("Executive (ผู้บริหาร)", "ROLE_EXEC", "เข้าถึง Executive Dashboard เพื่อดูภาพรวมความคืบหน้าของทุกโครงการ, วิเคราะห์แนวโน้มต้นทุน, Budget Burn Rate, กำไรขาดทุน (P&L) และดัชนีความล่าช้า (Drift)")
    ]
    
    for row_idx, (r_name, r_code, r_desc) in enumerate(roles_data, start=1):
        row = role_tbl.rows[row_idx]
        bg = "F9FBFD" if row_idx % 2 == 1 else "FFFFFF"
        for c_idx, text in enumerate([r_name, r_code, r_desc]):
            cell = row.cells[c_idx]
            cell.width = role_widths[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, 80, 80, 100, 100)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = 'TH Sarabun New'
            r.font.size = Pt(12)
            if c_idx == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(0x0F, 0x3B, 0x68)
            else:
                r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 2. ฟีเจอร์เด่นเพื่อการควบคุมคุณภาพ (Key Quality Control Features)
    add_heading_styled(doc, "2. เสาหลักฟีเจอร์ควบคุมคุณภาพ (Key Quality Control Features)", 1)
    
    features = [
        ("GPS & Camera Proof of Work:", " ควบคุมการปฏิบัติงานด้วยการตรวจสอบพิกัดตำแหน่งจริงผ่านดาวเทียม (GPS Geofencing) และการบังคับถ่ายภาพผ่าน HTML5 Camera API แบบสด เพื่อป้องกันการลงเวลาผิดสถานที่หรือการส่งภาพย้อนหลัง"),
        ("Gantt Baseline & Drift Variance Engine:", " ระบบล็อกแผนแม่บทเริ่มต้น (Frozen Baseline) เพื่อเปรียบเทียบกับความคืบหน้างานจริง ทำให้ผู้บริหารและ PM วิเคราะห์จำนวนวันที่ล่าช้า (Delay) และงบประมาณบานปลายได้ทันท่วงที"),
        ("Agile Kanban Board Integration:", " บริหารงานแบบ Task-based ที่เชื่อมโยงกับ Timesheet ทีมช่างสามารถลากการ์ดงานผ่านสถานะ To Do, In Progress, จนถึง Done พร้อมคำนวณต้นทุนค่าแรงตามชั่วโมงทำงานจริงอัตโนมัติ"),
        ("Real-time Executive Dashboard & P&L:", " สรุปตัวชี้วัดสำคัญระดับองค์กร เช่น Budget Burn Rate, Cash Flow, และจัดกลุ่มโครงการที่มีความเสี่ยงล่าช้า เพื่อการตัดสินใจเชิงกลยุทธ์ได้อย่างแม่นยำ")
    ]
    for title, desc in features:
        add_paragraph_styled(doc, text=desc, bold_prefix=title, space_after=4)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 3. รายละเอียดขั้นตอนการทำงาน 12 ขั้นตอน (Detailed 12 Steps Specification)
    add_heading_styled(doc, "3. ข้อกำหนดและขั้นตอนการดำเนินงาน 12 ลำดับ (The 12-Step Project Pipeline Specification)", 1)
    add_paragraph_styled(doc, "วงจรการบริหารโครงการในระบบ PMT แบ่งออกเป็น 3 ช่วงหลัก (3 Phases) และ 12 ขั้นตอนปฏิบัติการ โดยมีรายละเอียดข้อกำหนดทางเทคนิคและกระบวนการทำงานดังนี้:")

    steps_data = [
        # Phase 1
        {
            "num": 1,
            "phase": "1. Pre-Construction",
            "name": "Lead & Requirement Gathering (บันทึกข้อมูลลูกค้าและความต้องการ)",
            "obj": "บันทึกและคัดกรองข้อมูลลูกค้าเป้าหมาย (Customer Leads), ช่องทางการติดต่อ, ประเภทงาน (Renovate/Built-in/Installer), ขอบเขตงาน และกรอบงบประมาณเบื้องต้นของลูกค้า",
            "actor": "Responsible: ฝ่ายขาย / PM | Accountable: PM | Informed: Admin",
            "inputs": "ชื่อลูกค้า, ข้อมูลติดต่อ (เบอร์โทร/LINE/Email), ที่ตั้งหน้างาน (Location Pin), งบประมาณโดยประมาณ, ประเภทบริการ",
            "process": "1. เจ้าหน้าที่ฝ่ายขายหรือ PM รับข้อมูลความต้องการผ่านช่องทางต่างๆ เข้าสู่ระบบ\n2. บันทึก Lead ใหม่ พร้อมปักหมุดพิกัด Google Maps และระบุระดับความสำคัญ (Priority)\n3. ระบบจัดเก็บและส่ง Notification แจ้งเตือน PM หรือทีมงานในโซนพื้นที่รับผิดชอบ",
            "sa": "Entity `leads` จัดเก็บข้อมูลพร้อม Auto-generating Lead Code (เช่น LD-2026-0001) และทำ Data Validation เบอร์โทรและพิกัด GPS",
            "output": "บันทึก Lead ในระบบ สถานะ 'New Lead' และแสดงบนกระดาน Leads Management",
            "sow": "ระบบต้องสามารถบันทึกและค้นหาข้อมูล Lead ได้แบบเรียลไทม์ พร้อมแสดงตำแหน่งบนแผนที่ได้อย่างถูกต้อง"
        },
        {
            "num": 2,
            "phase": "1. Pre-Construction",
            "name": "Survey Booking (การนัดหมายสำรวจหน้างาน)",
            "obj": "จัดตารางเวลาและนัดหมายทีมช่างสำรวจพื้นที่หน้างานจริง พร้อมประสานงานยืนยันวัน-เวลากับลูกค้า เพื่อป้องกันการจัดคิวซ้อน (Double Booking)",
            "actor": "Responsible: PM / เจ้าหน้าที่จัดคิว | Accountable: PM | Informed: ช่างสำรวจ, ลูกค้า",
            "inputs": "Lead ID ที่ได้รับการคัดกรองแล้ว, วันที่และช่วงเวลาที่ลูกค้านัดหมาย, ทีมช่างสำรวจที่ว่าง",
            "process": "1. PM ตรวจสอบปฏิทินความพร้อมของช่างสำรวจ (Booking Calendar)\n2. กำหนดวัน-เวลา และมอบหมายช่างสำรวจเข้าพื้นที่\n3. ระบบส่งการแจ้งเตือนยืนยันการนัดหมายไปยังช่างสำรวจและส่งข้อความยืนยันแก่ลูกค้า",
            "sa": "Entity `surveys` พร้อมระบบ Conflict Detection ตรวจสอบไม่ให้มีการจองคิวช่างคนเดียวกันในเวลาซ้อนทับกัน และอัปเดตสถานะ Lead เป็น 'Survey Scheduled'",
            "output": "นัดหมายสำรวจถูกบันทึกลงใน Calendar กลาง และแสดงในแอปพลิเคชันของช่างสำรวจ",
            "sow": "มีระบบตรวจสอบคิวงานซ้อน และสามารถส่งข้อความแจ้งเตือนนัดหมายไปยังผู้เกี่ยวข้องได้โดยอัตโนมัติ"
        },
        {
            "num": 3,
            "phase": "1. Pre-Construction",
            "name": "Survey QC Inspection (การตรวจสำรวจและบันทึกสภาพเดิม)",
            "obj": "ตรวจสอบสภาพพื้นที่เดิม ขนาดพื้นที่จริง จุดเสี่ยง และอุปสรรคหน้างาน พร้อมบันทึกภาพถ่ายหลักฐานก่อนเริ่มงาน (Pre-work Evidence) ป้องกันข้อพิพาทความเสียหายเดิม",
            "actor": "Responsible: ช่างสำรวจ / Inspector | Accountable: PM | Consulted: ลูกค้า",
            "inputs": "รายการนัดหมายสำรวจ, แบบฟอร์ม Digital Survey Checklist, อุปกรณ์มือถือสำหรับถ่ายภาพและวัดขนาด",
            "process": "1. ช่างสำรวจเดินทางถึงหน้างาน ทำการเปิดระบบบนมือถือและ Check-in\n2. กรอกข้อมูลวัดขนาดพื้นที่ กว้างxยาวxสูง และตรวจสอบสภาพพื้นผิว น้ำ ไฟฟ้า ทางเข้า-ออก\n3. บันทึกภาพถ่ายสภาพหน้างานจริงผ่านระบบ (บังคับเปิดกล้องสด)\n4. ลูกค้าหรือตัวแทนหน้างานเซ็นชื่อรับทราบผลการสำรวจผ่านหน้าจออุปกรณ์ดิจิทัล",
            "sa": "เรียกใช้ Geolocation API เพื่อตรวจสอบพิกัด และ HTML5 Camera API สำหรับบันทึกภาพพร้อม Timestamp / GPS Metadata และแปลงข้อมูลเป็นรายงาน Survey Report",
            "output": "รายงานผลการสำรวจ (Digital Survey Report) พร้อมภาพถ่ายสภาพเดิมและลายเซ็นรับทราบ สถานะ 'Survey Completed'",
            "sow": "ระบบต้องบันทึกภาพถ่ายสภาพเดิมพร้อมพิกัด GPS และจัดเก็บเอกสารผลสำรวจเป็นหลักฐานอ้างอิงได้สมบูรณ์"
        },
        {
            "num": 4,
            "phase": "1. Pre-Construction",
            "name": "2D/3D Design & Space Planning (การจัดทำแบบแปลนและภาพจำลอง)",
            "obj": "จัดทำและอัปโหลดแบบแปลน 2D (Layout/Floor Plan) และภาพจำลองเสมือนจริง 3D (3D Rendering) เพื่อนำเสนอแนวคิดให้ลูกค้าเห็นภาพและอนุมัติแบบก่อนประเมินราคา",
            "actor": "Responsible: สถาปนิก / ดีไซเนอร์ | Accountable: PM | Consulted: ลูกค้า",
            "inputs": "ข้อมูลและรูปถ่ายจากผลสำรวจ Survey Report, ความต้องการสไตล์ของลูกค้า",
            "process": "1. ดีไซเนอร์ดึงข้อมูลพื้นที่และขนาดจริงไปดำเนินการเขียนแบบ 2D และทำภาพ 3D Perspective\n2. อัปโหลดไฟล์แบบและเอกสารประกอบเข้าระบบโครงการ\n3. PM นำเสนอแบบแก่ลูกค้าผ่านระบบหรือการประชุมเพื่อขอความเห็นชอบ\n4. ลูกค้าอนุมัติแบบ ระบบบันทึกประวัติการอนุมัติ (Design Approval)",
            "sa": "Entity `project_designs` รองรับ Version Control (Revision A, B, C) ป้องกันการนำแบบผิดเวอร์ชันไปใช้งานหน้างาน และรองรับการดูไฟล์ภาพ/PDF ผ่าน Web Viewer",
            "output": "ชุดแบบแปลน 2D/3D ที่ผ่านการอนุมัติจากลูกค้า และสถานะโครงการเปลี่ยนเป็น 'Design Approved'",
            "sow": "ระบบต้องมีระบบควบคุมเวอร์ชันของแบบแปลน และบันทึกประวัติการอนุมัติของลูกค้าได้อย่างโปร่งใส"
        },
        
        # Phase 2
        {
            "num": 5,
            "phase": "2. Construction (Budget & Preparation)",
            "name": "BOQ Cost Estimation (การถอดแบบราคากลางและประมาณการต้นทุน)",
            "obj": "คำนวณต้นทุนค่าวัสดุ ค่าแรง และค่าดำเนินการ โดยดึงราคาจากฐานข้อมูลราคากลางมาตรฐาน (Service Price Book) เพื่อความแม่นยำและรักษาอัตรากำไรเป้าหมาย (Target Gross Margin)",
            "actor": "Responsible: PM / QS (Quantity Surveyor) | Accountable: PM / ผู้บริหาร",
            "inputs": "แบบแปลนที่ได้รับอนุมัติ, ฐานข้อมูล Service Price Book Master Data, ปริมาณงาน (Quantities)",
            "process": "1. PM/QS สร้างเอกสาร BOQ ผูกกับโครงการที่ได้รับอนุมัติแบบ\n2. เลือกรายการวัสดุและบริการจาก Service Price Book (ระบบดึง Cost Rate และ Selling Rate อัตโนมัติ)\n3. ระบุปริมาณงาน ระบบทำการคำนวณต้นทุนรวมและราคาเสนอขาย\n4. ตรวจสอบ Gross Margin % และตั้งค่าเป็นงบประมาณแม่บท (Baseline Budget)",
            "sa": "Entity `boq_items`, `service_price_book` คำนวณ Material Cost + Labor Cost + Overhead + Margin = Total Price พร้อมล็อกสูตรป้องกันความผิดพลาด",
            "output": "เอกสาร BOQ ฉบับสมบูรณ์ และงบประมาณต้นทุนเป้าหมาย (Baseline Cost Budget)",
            "sow": "การประเมินราคาต้องอ้างอิงจากฐานข้อมูลราคากลาง และสามารถคำนวณแยกหมวดหมู่วัสดุ-ค่าแรงได้อย่างถูกต้องแม่นยำ"
        },
        {
            "num": 6,
            "phase": "2. Construction (Budget & Preparation)",
            "name": "Quotation Approval & Contracting (การเสนอราคาและเซ็นอนุมัติสัญญา)",
            "obj": "แปลงข้อมูล BOQ เป็นใบเสนอราคาทางการ (Quotation) เสนอต่อลูกค้า และดำเนินการลงนามเซ็นสัญญาข้อตกลงการจ้างผ่านระบบ",
            "actor": "Responsible: PM / ฝ่ายขาย | Accountable: PM | Approver: ผู้มีอำนาจลงนาม / ลูกค้า",
            "inputs": "เอกสาร BOQ ที่ผ่านการตรวจสอบ, เงื่อนไขการแบ่งจ่ายงวดงาน (Payment Milestones)",
            "process": "1. ระบบแปลง BOQ เป็นใบเสนอราคา พร้อมระบุเงื่อนไขการรับประกันและงวดการชำระเงิน\n2. ส่งใบเสนอราคาให้ลูกค้าพิจารณาทางดิจิทัล\n3. ลูกค้าลงนามอนุมัติ (E-Signature หรือแนบเอกสาร PO สั่งจ้าง)\n4. PM ยืนยันการอนุมัติในระบบ ระบบจะแปลง Lead ไปเป็น 'Active Project' พร้อมสร้างแผนงาน Gantt และ Kanban อัตโนมัติ",
            "sa": "Entity `quotations`, `projects`, `contracts` ทำการแปลงสถานะ Lead สู่ Project Instance และดึง Milestone Template มาสร้าง WBS อัตโนมัติ",
            "output": "ใบเสนอราคาและสัญญาที่ได้รับอนุมัติ และโครงการใหม่ถูกสร้างขึ้นในระบบ สถานะ 'Quotation Approved'",
            "sow": "ระบบต้องสร้างเอกสารใบเสนอราคาพร้อมเลขรันนิ่ง และแปลงโครงการเป็น Active Project พร้อมตารางงาน Gantt ทันทีที่อนุมัติ"
        },
        {
            "num": 7,
            "phase": "2. Construction (Budget & Preparation)",
            "name": "Down Payment & Financial Entry (การบันทึกชำระเงินมัดจำงวดแรก)",
            "obj": "บันทึกการรับชำระเงินมัดจำงวดแรก (Down Payment) ก่อนการปล่อยงานและจัดสรรทรัพยากรเข้าพื้นที่จริง เพื่อป้องกันความเสี่ยงทางการเงินขององค์กร",
            "actor": "Responsible: ฝ่ายการเงิน / บัญชี | Accountable: PM | Informed: ลูกค้า",
            "inputs": "ใบเสนอราคาที่ได้รับอนุมัติ, หลักฐานการโอนเงิน/สลิปของลูกค้า (Payment Slip)",
            "process": "1. ลูกค้าชำระเงินมัดจำงวดแรกตามเงื่อนไขสัญญา\n2. ฝ่ายการเงินตรวจสอบยอดเงิน และบันทึกการรับชำระในระบบพร้อมแนบสลิป\n3. ระบบอัปเดตสถานะทางการเงินของงวดงานที่ 1 เป็น 'Paid'\n4. ระบบปลดล็อกสิทธิ์ให้โครงการสามารถเริ่มขั้นตอน Site Check-In เข้าพื้นที่ได้",
            "sa": "Entity `project_payments`, `invoices` บันทึกกระแสเงินสดรับเข้าโครงการ (Cash Inflow) และส่งสัญญาณ Authorization ไปยังส่วน Project Execution",
            "output": "ใบเสร็จรับเงิน/ใบสำคัญรับ และสถานะโครงการเปลี่ยนเป็น 'Ready for Site Entry' (พร้อมเข้าหน้างาน)",
            "sow": "ระบบต้องมีระบบควบคุมความปลอดภัย ไม่อนุญาตให้เริ่มขั้นตอนเข้าหน้างานจนกว่าจะมีการยืนยันชำระเงินมัดจำเรียบร้อย"
        },
        {
            "num": 8,
            "phase": "2. Construction (Site Launch)",
            "name": "Site Check-In & Proof of Presence (การเข้าพื้นที่หน้างานและเช็คอิน GPS)",
            "obj": "ควบคุมและยืนยันการเข้าพื้นที่หน้างานจริงของทีมช่าง ด้วยระบบระบุพิกัดดาวเทียม (GPS Geofencing) และภาพถ่ายยืนยันตัวตน เพื่อป้องกันปัญหาการรายงานเท็จ",
            "actor": "Responsible: หัวหน้าช่าง / ทีมปฏิบัติการ | Accountable: PM | Informed: ลูกค้า",
            "inputs": "แอปพลิเคชันมือถือ, พิกัด GPS หน้างาน, กล้องมือถือสำหรับถ่ายภาพยืนยันตัวตน",
            "process": "1. ทีมช่างเดินทางถึงสถานที่ก่อสร้าง เปิดระบบบนมือถือและเลือกโครงการ\n2. กดปุ่ม 'Site Check-In' ระบบตรวจสอบพิกัดปัจจุบันเทียบกับพิกัดโครงการ (รัศมี Geofence 200-500 ม.)\n3. ช่างถ่ายภาพตนเองพร้อมบรรยากาศหน้างานผ่านกล้องสด\n4. กดยืนยัน ระบบบันทึกเวลาเข้างานและแจ้งเตือนบนหน้าจอ Dashboard ของ PM ทันที",
            "sa": "คำนวณระยะทางด้วย Haversine Formula เทียบกับพิกัดโครงการ บังคับถ่ายภาพสด (`capture=camera`) และบันทึกลง Entity `site_checkins`",
            "output": "บันทึกประวัติการเช็คอินพร้อมพิกัด เวลา และรูปถ่ายแสดงบน Live Dashboard ของโครงการ",
            "sow": "ระบบต้องตรวจสอบพิกัด GPS จริงและป้องกันการนำภาพเก่ามาใช้ โดยมีระบบเตือนเมื่ออยู่นอกพื้นที่ที่กำหนด"
        },
        
        # Phase 3
        {
            "num": 9,
            "phase": "3. Completion & Review",
            "name": "Execution & Daily Timesheet (การดำเนินงานและบันทึกเวลาทำงานรายวัน)",
            "obj": "ดำเนินงานก่อสร้าง/ติดตั้งตามแผน และบันทึกเวลาการทำงาน (Timesheet) พร้อมอัปเดตสถานะงานบน Agile Kanban Board เพื่อใช้คำนวณต้นทุนค่าแรงจริง (Labor Cost Accrual)",
            "actor": "Responsible: ทีมช่าง / พนักงานปฏิบัติการ | Accountable: PM | Consulted: หัวหน้าช่าง",
            "inputs": "การ์ดงานบนบอร์ด Kanban, ชั่วโมงทำงานจริง (ชั่วโมงปกติ / OT), ภาพถ่ายผลงานประจำวัน",
            "process": "1. ช่างปฏิบัติงานตาม Task และเลื่อนสถานะการ์ดบน Kanban Board (To Do -> In Progress -> Done)\n2. สิ้นสุดวันทำงาน ช่างบันทึกชั่วโมงการทำงานและแนบภาพถ่ายผลงานที่ทำเสร็จในวันนั้น\n3. ส่งใบลงเวลา (Submit Timesheet) ไปยัง PM เพื่อรอการตรวจรับรอง",
            "sa": "Entity `timesheets`, `tasks`, `kanban_cards` คำนวณ Actual Labor Cost โดยนำชั่วโมงคูณกับ User Cost Rate และสะสมเข้า Project Cost Real-time",
            "output": "บันทึก Timesheet รายวัน และยอดต้นทุนค่าแรงสะสมถูกอัปเดตในระบบ",
            "sow": "ระบบต้องรองรับการบันทึก Timesheet ผูกกับ Task บน Kanban และคำนวณต้นทุนค่าแรงจริงสะสมรายโครงการได้โดยอัตโนมัติ"
        },
        {
            "num": 10,
            "phase": "3. Completion & Review",
            "name": "Daily QC & Progress Update (การตรวจคุณภาพรายวันและอัปเดตความคืบหน้า)",
            "obj": "ตรวจสอบคุณภาพงานตามเกณฑ์ Milestone, อนุมัติใบลงเวลา (Timesheet Approval), และอัปเดต % ความคืบหน้าจริง เทียบกับแผนแม่บทเพื่อตรวจจับความล่าช้า (Gantt Baseline Drift)",
            "actor": "Responsible: PM / Site Engineer / QC Inspector | Accountable: PM",
            "inputs": "รายการ Timesheet ที่ส่งเข้ามา, ภาพถ่ายผลงานรายวัน, รายการ Checklist ตรวจสอบคุณภาพ",
            "process": "1. PM/วิศวกรเข้าตรวจงานหน้างานหรือตรวจผ่านระบบดิจิทัล\n2. พิจารณาอนุมัติ (Approve) หรือส่งกลับแก้ไข (Reject) รายการ Timesheet พร้อมระบุเหตุผล\n3. ตรวจเช็คเกณฑ์คุณภาพ QC Checklist ประจำ Milestone\n4. อัปเดต % ความคืบหน้าของโครงการ ระบบจะคำนวณความคลาดเคลื่อนและจำนวนวันที่ล่าช้า (Drift Days)",
            "sa": "คำนวณเปรียบเทียบ Baseline Timeline กับ Live Timeline เพื่อแสดงผลบน Gantt Chart Drift Analysis และบันทึกผล QC ลง Entity `qc_inspections`",
            "output": "Timesheet ได้รับการอนุมัติ, รายงานผล QC ประจำวัน, และกราฟความคืบหน้าอัปเดตตามจริง",
            "sow": "ระบบต้องมีระบบ Workflow อนุมัติ Timesheet โดย PM และมีกลไกเปรียบเทียบ Gantt Baseline vs Actual Drift แสดงผลแบบ Real-time"
        },
        {
            "num": 11,
            "phase": "3. Completion & Review",
            "name": "Final QC & Customer Handover (การตรวจรับงานงวดสุดท้ายและส่งมอบผลงาน)",
            "obj": "ตรวจสอบคุณภาพขั้นสุดท้าย (Defect Inspection), แก้ไขรายการเก็บงาน (Punch List), และส่งมอบผลงานที่สมบูรณ์ให้แก่ลูกค้าพร้อมลงนามรับมอบงานอย่างเป็นทางการ",
            "actor": "Responsible: PM / QC Lead | Accountable: PM | Approver: ลูกค้า",
            "inputs": "ผลงานที่เสร็จสิ้น 100%, รายการ Defect Checklist, เอกสารใบส่งมอบงาน (Handover Certificate)",
            "process": "1. PM นัดหมายลูกค้าตรวจสอบผลงานงวดสุดท้ายร่วมกันหน้างาน\n2. บันทึกรายการ Defect ที่ต้องเก็บงาน (ถ้ามี) เข้าสู่ระบบ และมอบหมายช่างเข้าแก้ไข\n3. เมื่อแก้ไขครบถ้วน ลูกค้าลงนามตรวจรับมอบงานดิจิทัล (Digital Handover Sign-off)\n4. บันทึกการรับชำระเงินงวดสุดท้าย และออกเอกสารเริ่มนับระยะเวลารับประกันผลงาน (Warranty Period)",
            "sa": "Entity `handover_certificates`, `defect_items` สร้างเอกสาร PDF ใบส่งมอบงานพร้อมลายเซ็นดิจิทัล และอัปเดตสถานะโครงการเป็น 'Delivered / Pending Closeout'",
            "output": "ใบส่งมอบงานที่ลงนามสมบูรณ์, รายการ Defect ปิดครบถ้วน, และเอกสารเริ่มประกันผลงาน",
            "sow": "ระบบต้องมีฟังก์ชันบันทึกติดตามการแก้ Defect จนครบถ้วน และสร้างใบรับมอบงานพร้อมลายเซ็นดิจิทัลของลูกค้าได้"
        },
        {
            "num": 12,
            "phase": "3. Completion & Review",
            "name": "Project Closeout & Evaluation (การปิดโครงการและประเมินผลกำไรขาดทุน)",
            "obj": "ปิดโครงการอย่างเป็นทางการ สรุปงบกำไร-ขาดทุนจริง (P&L: Revenue vs Actual Material + Labor Cost), และประเมินความพึงพอใจของลูกค้า (CSAT) เพื่อเก็บสถิติองค์กร",
            "actor": "Responsible: PM / ฝ่ายการเงิน | Accountable: ผู้บริหาร (Executive) | Informed: ทุกฝ่าย",
            "inputs": "ยอดรายรับทั้งหมดที่ชำระครบถ้วน, ยอดค่าใช้จ่ายจริงทั้งหมด (วัสดุ + ค่าแรง Timesheet + อื่นๆ), แบบประเมิน CSAT",
            "process": "1. ฝ่ายการเงินและ PM ตรวจสอบความถูกต้องของรายรับ-รายจ่ายทั้งหมดในโครงการ\n2. ระบบประมวลผลคำนวณผลกำไรขั้นต้นสุทธิ (Gross Profit Margin %) และสรุปความล่าช้าจริง\n3. ส่งแบบสอบถามความพึงพอใจดิจิทัลให้ลูกค้าประเมินคะแนนการบริการ\n4. PM กดปุ่ม 'Close Project' ระบบทำการล็อกโครงการเป็นสถานะอ่านอย่างเดียว (Archive)",
            "sa": "Entity `projects` ปรับสถานะเป็น 'Closed' คำนวณ Financial Metrics ส่งเข้า Executive Data Lake และส่งข้อมูลสรุปไปยัง Executive Dashboard",
            "output": "รายงานสรุปปิดโครงการ (Project Closeout P&L Report) และคะแนนความพึงพอใจลูกค้า",
            "sow": "ระบบต้องสรุปผลกำไร-ขาดทุนจริงเปรียบเทียบกับงบประมาณเริ่มต้นได้โดยอัตโนมัติ และมีระบบล็อกข้อมูลเมื่อปิดโครงการสมบูรณ์"
        }
    ]

    current_phase = ""
    for step in steps_data:
        if step["phase"] != current_phase:
            current_phase = step["phase"]
            add_heading_styled(doc, f"ช่วงการดำเนินงาน: {current_phase}", 2)
            
        add_heading_styled(doc, f"ขั้นตอนที่ {step['num']}: {step['name']}", 3)
        
        # Step Table Specification
        tbl = doc.add_table(rows=7, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl, "B0C4DE")
        s_widths = [Inches(1.8), Inches(5.0)]
        
        fields = [
            ("วัตถุประสงค์ (Objectives)", step["obj"]),
            ("ผู้รับผิดชอบ (RACI Roles)", step["actor"]),
            ("ข้อมูลนำเข้า (Inputs / Prerequisites)", step["inputs"]),
            ("ขั้นตอนการดำเนินงาน (Process Flow)", step["process"]),
            ("การทำงานระบบ (SA & System Mechanics)", step["sa"]),
            ("ผลลัพธ์ที่ได้ (Outputs / Deliverables)", step["output"]),
            ("เกณฑ์ตรวจรับงาน (SOW Acceptance Criteria)", step["sow"])
        ]
        
        for r_idx, (f_label, f_val) in enumerate(fields):
            row = tbl.rows[r_idx]
            c0, c1 = row.cells[0], row.cells[1]
            c0.width, c1.width = s_widths[0], s_widths[1]
            set_cell_background(c0, "F0F4F8")
            set_cell_background(c1, "FFFFFF" if r_idx % 2 == 0 else "FAFCFE")
            set_cell_margins(c0, 60, 60, 100, 100)
            set_cell_margins(c1, 60, 60, 100, 100)
            
            p0 = c0.paragraphs[0]
            p0.paragraph_format.space_before = Pt(0)
            p0.paragraph_format.space_after = Pt(0)
            r0 = p0.add_run(f_label)
            r0.font.name = 'TH Sarabun New'
            r0.font.size = Pt(11.5)
            r0.bold = True
            r0.font.color.rgb = RGBColor(0x0F, 0x3B, 0x68)
            
            p1 = c1.paragraphs[0]
            p1.paragraph_format.space_before = Pt(0)
            p1.paragraph_format.space_after = Pt(0)
            r1 = p1.add_run(f_val)
            r1.font.name = 'TH Sarabun New'
            r1.font.size = Pt(11.5)
            r1.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 4. ตารางเมทริกซ์ความรับผิดชอบ (RACI Responsibility Matrix)
    add_heading_styled(doc, "4. ตารางสรุปความรับผิดชอบในแต่ละขั้นตอน (RACI Matrix)", 1)
    add_paragraph_styled(doc, "คำอธิบาย: R = Responsible (ผู้ปฏิบัติงาน), A = Accountable (ผู้รับผิดชอบผลงาน/อนุมัติ), C = Consulted (ผู้ให้คำปรึกษา), I = Informed (ผู้รับทราบข้อมูล)")

    raci_tbl = doc.add_table(rows=13, cols=6)
    raci_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(raci_tbl, "CCCCCC")
    raci_widths = [Inches(0.6), Inches(2.6), Inches(0.9), Inches(0.9), Inches(0.9), Inches(0.9)]
    
    raci_headers = ["ลำดับ", "ขั้นตอนการทำงาน (Pipeline Step)", "Admin", "PM", "Technician", "Executive"]
    for i, h in enumerate(raci_headers):
        cell = raci_tbl.rows[0].cells[i]
        cell.width = raci_widths[i]
        set_cell_background(cell, "0F3B68")
        set_cell_margins(cell, 80, 80, 80, 80)
        p = cell.paragraphs[0]
        if i >= 2:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = 'TH Sarabun New'
        r.font.size = Pt(12)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    raci_matrix = [
        ("1", "Lead & Requirement Gathering", "I", "R/A", "-", "I"),
        ("2", "Survey Booking", "I", "R/A", "C", "I"),
        ("3", "Survey QC Inspection", "-", "A", "R", "I"),
        ("4", "2D/3D Design & Planning", "-", "A", "-", "I"),
        ("5", "BOQ Cost Estimation", "C", "R/A", "-", "I"),
        ("6", "Quotation Approval", "I", "R/A", "-", "A/I"),
        ("7", "Down Payment Entry", "I", "R/A", "-", "I"),
        ("8", "Site Check-In (GPS & Photo)", "-", "A", "R", "I"),
        ("9", "Execution & Daily Timesheet", "-", "A", "R", "I"),
        ("10", "Daily QC & Progress Update", "-", "R/A", "C", "I"),
        ("11", "Final QC & Handover", "-", "R/A", "C", "I"),
        ("12", "Project Closeout & P&L", "I", "R", "-", "A")
    ]
    
    for row_idx, r_data in enumerate(raci_matrix, start=1):
        row = raci_tbl.rows[row_idx]
        bg = "F9FBFD" if row_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(r_data):
            cell = row.cells[c_idx]
            cell.width = raci_widths[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, 60, 60, 80, 80)
            p = cell.paragraphs[0]
            if c_idx >= 2 or c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            r.font.name = 'TH Sarabun New'
            r.font.size = Pt(11.5)
            if "R" in val or "A" in val:
                r.bold = True
                r.font.color.rgb = RGBColor(0x0F, 0x3B, 0x68)
            else:
                r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # 5. การลงนามรับรองเอกสารขอบเขตงาน (Sign-off Section)
    add_heading_styled(doc, "5. การลงนามรับทราบและเห็นชอบขอบเขตงาน (SOW Acceptance Sign-Off)", 1)
    add_paragraph_styled(doc, "เอกสารข้อกำหนดขอบเขตงานระบบ PMT ฉบับนี้ ได้รับการตรวจสอบและเห็นชอบจากผู้มีส่วนเกี่ยวข้อง เพื่อใช้เป็นเกณฑ์มาตรฐานในการพัฒนา ตรวจรับมอบงาน และควบคุมคุณภาพโครงการ:")

    sign_tbl = doc.add_table(rows=2, cols=2)
    sign_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(sign_tbl, "CCCCCC")
    sign_widths = [Inches(3.4), Inches(3.4)]
    
    sign_boxes = [
        ("สำหรับ: ผู้ว่าจ้าง / ผู้บริหารโครงการ (Client / Management)", "\n\n\nลงชื่อ: ........................................................\n( ........................................................ )\nตำแหน่ง: ....................................................\nวันที่: ......./......./............"),
        ("สำหรับ: ผู้จัดการโครงการและสถาปนิกพัฒนาระบบ (PM & Lead SA)", "\n\n\nลงชื่อ: ........................................................\n( ........................................................ )\nตำแหน่ง: ....................................................\nวันที่: ......./......./............")
    ]
    
    for c_idx, (s_title, s_content) in enumerate(sign_boxes):
        cell = sign_tbl.rows[0].cells[c_idx]
        cell.width = sign_widths[c_idx]
        set_cell_background(cell, "F0F4F8")
        set_cell_margins(cell, 80, 80, 100, 100)
        p = cell.paragraphs[0]
        r = p.add_run(s_title)
        r.font.name = 'TH Sarabun New'
        r.font.size = Pt(11.5)
        r.bold = True
        r.font.color.rgb = RGBColor(0x0F, 0x3B, 0x68)
        
        cell2 = sign_tbl.rows[1].cells[c_idx]
        cell2.width = sign_widths[c_idx]
        set_cell_background(cell2, "FFFFFF")
        set_cell_margins(cell2, 100, 100, 100, 100)
        p2 = cell2.paragraphs[0]
        r2 = p2.add_run(s_content)
        r2.font.name = 'TH Sarabun New'
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.save(output_path)
    print(f"Successfully generated: {output_path}")

if __name__ == "__main__":
    out_file = r"c:\atgv\vbooking\PMT_12_Steps_SOW_Specification.docx"
    img_file = r"C:\Users\isara\.gemini\antigravity\brain\a47b4d83-2349-4a12-bd2a-ed2291707780\.user_uploaded\media_1787020048617.jpg"
    build_docx(out_file, img_file)
